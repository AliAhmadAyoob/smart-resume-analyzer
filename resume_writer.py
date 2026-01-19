from docx import Document

def build_updated_resume(parsed):
    doc = Document()

    # Name
    doc.add_heading(parsed.get("name", "Candidate"), level=0)

    # Contact
    contact = parsed.get("contact", [])
    if contact:
        doc.add_heading("Contact Information", level=1)
        for item in contact:
            doc.add_paragraph(str(item))

    # Skills
    skills = parsed.get("skills", [])
    if skills:
        doc.add_heading("Skills", level=1)
        for s in skills:
            doc.add_paragraph(f"• {s}")

    # Experience
    exp = parsed.get("experience", [])
    if exp:
        doc.add_heading("Experience", level=1)
        for e in exp:
            doc.add_paragraph(f"• {e}")

    # Projects
    proj = parsed.get("projects", [])
    if proj:
        doc.add_heading("Projects", level=1)
        for p in proj:
            doc.add_paragraph(f"• {p}")

    filename = "updated_resume.docx"
    doc.save(filename)
    return filename
